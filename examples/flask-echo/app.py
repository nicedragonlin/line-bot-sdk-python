# -*- coding: utf-8 -*-

#  Licensed under the Apache License, Version 2.0 (the "License"); you may
#  not use this file except in compliance with the License. You may obtain
#  a copy of the License at
#
#       https://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS, WITHOUT
#  WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied. See the
#  License for the specific language governing permissions and limitations
#  under the License.


import os
import sys
import yaml
import math
from docx.api import Document
import pandas as pd
from argparse import ArgumentParser
import requests
import gdown
import re

from flask import Flask, request, abort
from linebot import (
    WebhookParser
)
from linebot.v3.exceptions import (
    InvalidSignatureError
)
from linebot.v3.webhooks import (
    MessageEvent,
    TextMessageContent,
)
from linebot.v3.messaging import (
    Configuration,
    ApiClient,
    MessagingApi,
    ReplyMessageRequest,
    TextMessage
)

searve_people_num = 480


app = Flask(__name__)

# get channel_secret and channel_access_token from your environment variable
channel_secret = os.getenv('LINE_CHANNEL_SECRET', None)
channel_access_token = os.getenv('LINE_CHANNEL_ACCESS_TOKEN', None)
if channel_secret is None:
    print('Specify LINE_CHANNEL_SECRET as environment variable.')
    sys.exit(1)
if channel_access_token is None:
    print('Specify LINE_CHANNEL_ACCESS_TOKEN as environment variable.')
    sys.exit(1)

parser = WebhookParser(channel_secret)

configuration = Configuration(
    access_token=channel_access_token
)


@app.route("/callback", methods=['POST'])
def callback():
    signature = request.headers['X-Line-Signature']

    # get request body as text
    body = request.get_data(as_text=True)
    app.logger.info("Request body: " + body)

    # parse webhook body
    try:
        events = parser.parse(body, signature)
    except InvalidSignatureError:
        abort(400)

    # if event is MessageEvent and message is TextMessage, then echo text
    for event in events:
        
        print('[Debug 100]')
        #if not isinstance(event, MessageEvent):
        #    print('[Debug 101]')
        #    continue
        #if not isinstance(event.message, TextMessageContent):
        #    print('[Debug 102]')
        #    continue
        with ApiClient(configuration) as api_client:
            print('[Debug 103]')
            line_bot_api = MessagingApi(api_client)
            line_bot_api.reply_message_with_http_info(
                ReplyMessageRequest(
                    reply_token=event.reply_token,
                    #messages=[TextMessage(text=event.message.text)]
                    messages=[TextMessage(text=miranda_list_ingredient(event.message.text))]
                )
            )
            print('[Debug 105]')
            
    return 'OK'


if __name__ == "__main__":
    arg_parser = ArgumentParser(
        usage='Usage: python ' + __file__ + ' [--port <port>] [--help]'
    )
    arg_parser.add_argument('-p', '--port', type=int, default=8000, help='port')
    arg_parser.add_argument('-d', '--debug', default=False, help='debug')
    options = arg_parser.parse_args()

    app.run(debug=options.debug, port=options.port)


def miranda_list_ingredient(download_link):
    print('[Debug M000]')
    # Load ingredient
    with open('ingredient.yaml','r',encoding='utf-8') as f:
        ingredient_dic = yaml.safe_load(f)
        #print(ingredient_dic)
    print(ingredient_dic)
    print('[Debug M001]')

    print('[Debug M002-1]')

    url = 'https://docs.google.com/document/d/15fLN8GigE071EFcm095LJH3WJXSjE7UE/edit?usp=sharing&ouid=105440516119125989146&rtpof=true&sd=true'
    url = download_link
    file_id = url.split('/')[-2]
    prefix = 'https://drive.google.com/uc?/export=download&confirm=1&id='
    gdown.download(prefix+file_id,'menu_download.docx')
    
    #document = Document('menu.docx')
    print('[Debug M002-2]')
    document = Document('menu_download.docx')

    data = []
    keys = ('date','breakfast', 'dish', 'dessert' )
    print('[Debug M002]')
    
    for table in document.tables : 
        for i, row in enumerate(table.rows):
            text = [cell.text for cell in row.cells]
            print('[Debug M003]')
            print(text)
            # Remove duplicate cell (somehow it happends eventhough it seems normal in doc)
            duplicate_list = []
            for i in range(len(text)-1):
                if i==0:
                    continue
                elif text[i-1]==text[i]:
                    duplicate_list.append(i)
            for i in range(len(duplicate_list)-1,-1,-1):
                del text[duplicate_list[i]]
            print('[Debug M004]')
            print(text)
            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            print('[DEBUG row_data]')
            print(row_data)
            if '月' not in row_data['date']:
                print('[WARN] Found a date not including "月": '+row_data['date'])
                continue
            if len(row_data) !=4:
                print('[WARN] Found a row with non-4-len tuple')
                continue

            # breakfast ingredient is buy in one day before. Every Monday breakfast is 家樂氏玉米片/牛奶 and it doens't need to process
            row_data_date = { 'date': row_data['date'], 'dish': row_data['dish']+'、'+row_data['dessert']}
            data.append(row_data_date)
            if '家樂氏玉米片' not in row_data['breakfast']: #not Monday's breakfast
                data[len(data)-2]['dish'] += '、'+row_data['breakfast']

    
    print(data)
    
    to_buy_per_day = [] # [ dish0_info, dish1_info]; dish0_info: { date: , ingredient: {} }
    for date_data in data:
        if date_data['date']=='':
            continue
        print('Date: ' + date_data['date'])
        dish_list = re.split('、|/', date_data['dish'])
        date_dish_ingredient = dict()
        found_dish_list = []
        not_found_dish_list = []
        
        for search_dish in dish_list:
            search_dish = search_dish.strip()
            if search_dish in ingredient_dic.keys():
                print('\t'+'Found dish: '+search_dish)
                found_dish_list.append(search_dish)
                for ingredient,mount_info in ingredient_dic[search_dish].items():
                    num_unit = searve_people_num/mount_info['people']
                    print('\t\t'+ingredient+' Num_unit: ' + str(num_unit))
                    if ingredient in date_dish_ingredient.keys()  : # exist ingredient
                        date_dish_ingredient[ingredient]['value'] += mount_info['value']*num_unit
                    else:
                        date_dish_ingredient[ingredient] = {'value': mount_info['value']*num_unit, 'unit': mount_info['unit']}
            else:
                not_found_dish_list.append(search_dish)
        # round up for each ingredient
        for  ingredient in date_dish_ingredient.keys():
            if date_dish_ingredient[ingredient]['unit'] == '斤':
                date_dish_ingredient[ingredient]['value'] = math.ceil(date_dish_ingredient[ingredient]['value']*2.0)/2
            else:
                date_dish_ingredient[ingredient]['value'] = math.ceil(date_dish_ingredient[ingredient]['value'])
        date_dish_dict = {'date': date_data['date'], 'ingredient': date_dish_ingredient, 'found_dish': found_dish_list, 'not_found_dish': not_found_dish_list}
        to_buy_per_day.append(date_dish_dict)
    
    to_buy_per_day = tuple(to_buy_per_day)
    print(to_buy_per_day)

    # output result
    result = ''
    for date_info in to_buy_per_day:
        result += date_info['date']+'\n'
        for ingredient,mount_info in date_info['ingredient'].items():
            if '少許' in mount_info['unit'] or '適量' in mount_info['unit']:
                result += '\t' + ingredient + ' ' + mount_info['unit'] + '\n'
            else:
    result += '-----------\n'
    for date_info in to_buy_per_day:
        result += '\n'+date_info['date'] + '\n'
        result += '\t找到: '
        for dish in date_info['found_dish']:
            result += dish+' '
        result += '\n'
        result += '\t沒找到: '
        for dish in date_info['not_found_dish']:
            result += dish + ' '            result += '\t'+ingredient+' '+str(mount_info['value'])+mount_info['unit']+'\n'

    
    print(result)
    return result
